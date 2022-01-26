from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document('test.docx')
table = document.add_table(rows=2, cols=4)
table.style = document.styles['Table Grid']

tr = table.rows[0].cells
tr[0].paragraphs[0].add_run('A')
tr[2].paragraphs[0].add_run('B')
tr[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

table.cell(0, 0).merge(table.cell(1, 0))

document.save('new.docx')

