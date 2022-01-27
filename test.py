from PySide6.QtCore import QThread, QObject, Signal, Slot
from docx import Document, shared
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.shared import Cm, RGBColor
from docx.oxml import parse_xml
from module import open_api_pars

# dd dd dd
class TEST:
    def __init__(self):
        key = 'sfSPRX+xNEExRUqE4cdhNjBSk4uXIv8F1CfLen06hdPGn5cflLJqy/nxmh48uF8fvdGk68k6Z5jWsU1n6BeNPA=='
        binfo = {'sigungu': '11260', 'bjdong': '10100',
                 'bun': '0090', 'ji': '0027', 'dong': ''}

        # self.get_rooms_thread = brt.BuildingRegisterThread(binfo, ["전유부", "소유자"])
        # self.get_rooms_thread.start()
        # self.get_rooms_thread.threadEvent.workerThreadDone.connect(self.test01)
        self.get_rooms_thread = open_api_pars.TestThread()
        self.get_rooms_thread.start()
        self.get_rooms_thread.threadEvent.workerThreadDone.connect(self.test01)

    @Slot(object)
    def test01(self, data):
        print("값")

        # for i in column:
        #     print(column)
        # self.document = Document('매매 계약서.docx')
        # self.add_table('임대인')
        # self.add_table('임차인')
        # self.add_table2('중개사')
        #
        # tables = self.document.tables
        # tables[0].rows[10].cells[1].paragraphs[0].text = 'ssssss'
        #
        # out_file = 'New.docx'
        # self.document.save(out_file)
        # print('저장 완료')

    def add_table(self, position):
        if position == "임대인":
            position = "임 대 인"
        elif position == "임차인":
            position = "임 차 인"
        elif position == "매도인":
            position = "매 도 인"
        elif position == "매수인":
            position = "매 수 인"
        elif position == "공동명의인":
            position = "공  동\n명 의 인"

        rgb_lightblue = RGBColor(222, 234, 246)
        rgb_lightgray = RGBColor(242, 242, 242)
        rgb_darkgray = RGBColor(89, 89, 89)

        table = self.document.add_table(rows=2, cols=8)
        table.style = self.document.styles['Table Grid']
        table.style.font.name = '맑은 고딕(본문)'
        table.style.font.size = shared.Pt(9)
        table.allow_autofit = True

        line_0 = table.rows[0].cells
        line_1 = table.rows[1].cells

        line_0[0].text = position
        line_0[0].merge(line_1[0])
        line_0[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        line_0[1].text = '주    소'
        line_0[2].merge(line_0[6])

        p = table.cell(0, 7).paragraphs[0].add_run('(인)')
        p.font.color.rgb = rgb_darkgray
        line_0[7].merge(line_1[7])
        line_0[7].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        line_1[1].text = '주민등록번호'
        line_1[3].text = '연락처'
        line_1[5].text = '성　명'

        line_0[0].width = Cm(2)
        line_0[1].width = Cm(2.8)
        line_1[2].width = Cm(3.5)
        line_1[3].width = Cm(1.7)
        line_1[4].width = Cm(3.5)
        line_1[5].width = Cm(1.7)
        line_1[6].width = Cm(2.3)

        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                t = table.rows[r].cells[c].paragraphs[0].text
                if t != '':
                    table.rows[r].cells[c].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        self.cell_color(table, 0, 0, rgb_lightblue)
        self.cell_color(table, 0, 1, rgb_lightgray)
        self.cell_color(table, 1, 1, rgb_lightgray)
        self.cell_color(table, 1, 3, rgb_lightgray)
        self.cell_color(table, 1, 5, rgb_lightgray)

    def add_table2(self, position):
        if position == "중개사":
            position = "개  업\n중 개 사"
        elif position == "공동중개사":
            position = "공  동\n중 개 사"

        rgb_lightblue = RGBColor(222, 234, 246)
        rgb_lightgray = RGBColor(242, 242, 242)
        rgb_darkgray = RGBColor(89, 89, 89)
        rgb_gray = RGBColor(191, 191, 191)

        table = self.document.add_table(rows=3, cols=7)
        table.style = self.document.styles['Table Grid']
        table.style.font.name = '맑은 고딕(본문)'
        table.style.font.size = shared.Pt(9)
        table.allow_autofit = True

        line_0 = table.rows[0].cells
        line_1 = table.rows[1].cells
        line_2 = table.rows[2].cells

        line_0[0].text = position
        line_0[0].merge(line_2[0])
        line_0[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        line_0[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        line_0[1].text = '사무소 소재지'
        line_0[2].merge(line_0[6])

        line_1[1].text = '사무소 명칭'
        line_1[2].merge(line_1[4])

        line_1[5].text = '대 표 자 명'

        p = table.cell(1, 6).paragraphs[0].add_run('서명 및 날인  ')
        p2 = table.cell(1, 6).paragraphs[0].add_run('(인)')
        p.font.color.rgb = rgb_gray
        p2.font.color.rgb = rgb_darkgray

        line_2[1].text = '등 록 번 호'
        line_2[3].text = '연락처'
        line_2[5].text = '소속공인중개사'

        p = table.cell(2, 6).paragraphs[0].add_run('서명 및 날인  ')
        p2 = table.cell(2, 6).paragraphs[0].add_run('(인)')
        p.font.color.rgb = rgb_gray
        p2.font.color.rgb = rgb_darkgray

        line_0[0].width = Cm(2)
        line_0[1].width = Cm(2.8)
        line_1[5].width = Cm(2.8)
        line_2[2].width = Cm(3.5)
        line_2[3].width = Cm(1.7)
        line_2[4].width = Cm(3.5)

        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                t = table.rows[r].cells[c].paragraphs[0].text
                if t != '':
                    if c == 0:
                        self.cell_color(table, r, c, rgb_lightblue)
                        continue
                    elif c == 6:
                        continue
                    self.cell_color(table, r, c, rgb_lightgray)
                    table.rows[r].cells[c].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        # self.cell_color(table, 0, 0, rgb_lightblue)
        # self.cell_color(table, 0, 1, rgb_lightgray)
        # self.cell_color(table, 1, 1, rgb_lightgray)
        # self.cell_color(table, 1, 3, rgb_lightgray)
        # self.cell_color(table, 1, 5, rgb_lightgray)

    @staticmethod
    def cell_color(table, y, x, color_str):
        shading = locals()
        shading = parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=color_str))
        table.rows[y].cells[x]._tc.get_or_add_tcPr().append(shading)


TEST()
