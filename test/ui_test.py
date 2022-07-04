import re
from docx import Document

test_items = """
<html><head><meta name="qrichtext" content="1" /><style type="text/css">
p, li { white-space: pre-wrap; }
</style></head><body style=" font-family:'Gulim'; font-size:9pt; font-weight:400; font-style:normal;">
<p style=" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;">안녕<span style=" font-weight:600;">하</span>세요<span style=" font-size:14pt;">. 장</span>성<span style=" font-style:italic;">남입니다</span></p></body></html>
"""



class HTMLHelper(object):
    """ Translates some html into word runs. """

    def __init__(self):
        self.get_tags = re.compile("(<[a-z,A-Z]+>)(.*?)(</[a-z,A-Z]+>)")

    def html_to_run_map(self, html_fragment):
        """ breakes an html fragment into a run map """
        ptr = 0
        run_map = []
        for match in self.get_tags.finditer(html_fragment):
            if match.start() > ptr:
                text = html_fragment[ptr:match.start()]
                if len(text) > 0:
                    run_map.append((text, "plain_text"))
            run_map.append((match.group(2), match.group(1)))
            ptr = match.end()
        if ptr < len(html_fragment) - 1:
            run_map.append((html_fragment[ptr:], "plain_text"))
        return run_map

    def insert_runs_from_html_map(self, paragraph, run_map):
        """ inserts some runs into a paragraph object. """
        for run_item in run_map:
            run = paragraph.add_run(run_item[0])
            if run_item[1] == "<i>":
                run.italic = True


doc = Document()
html_helper = HTMLHelper()
run_map = html_helper.html_to_run_map(test_items)
par = doc.add_paragraph()
html_helper.insert_runs_from_html_map(par, run_map)
doc.save("test_run_mapping.docx")
