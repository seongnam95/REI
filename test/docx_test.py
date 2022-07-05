from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import *


class WritingContract:
    def __init__(self):
        self.document = Document()
        self.document.add_heading('Document Title', 3)
        self.grid_t_style = self.document.styles["Table Grid"]

        self.set_doc()
        self.set_content()

        self.document.save("table4.docx")

    def set_doc(self):
        sec = self.document.sections[0]

        sec.left_margin = Cm(1.27)
        sec.right_margin = Cm(1.27)
        sec.top_margin = Cm(1.27)
        sec.bottom_margin = Cm(1.27)

    def set_content(self):

        table = self.document.add_table(12, 7)
        table.style = self.grid_t_style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 1. 부동산의 표시
        self.set_merge(table, (0, 0), (0, 6))
        self.set_merge(table, (1, 1), (1, 6))
        self.set_merge(table, (2, 2), (2, 4))
        self.set_merge(table, (4, 1), (4, 4))

        table.cell(0, 0).text = '임대인과 임차인 쌍방은 아래 표시 부동산에 관하여 다음 계약내용과 같이 임대차계약을 체결한다.'
        table.cell(1, 0).text = '소 재 지'
        table.cell(2, 0).text = '토 　 지'
        table.cell(2, 1).text = '지 　 목'
        table.cell(2, 5).text = '면 　 적'
        table.cell(3, 0).text = '건 　 물'
        table.cell(3, 1).text = '구 　 조'
        table.cell(3, 3).text = '용 　 도'
        table.cell(3, 5).text = '면 　 적'
        table.cell(4, 0).text = '임대할부분'
        table.cell(4, 5).text = '면 　 적'

    @staticmethod
    def set_merge(t, start, end):
        t.cell(start[0], start[1]).merge(t.cell(end[0], end[1]))

    @staticmethod
    def set_cell(t, row, col):
        return t.rows[row].cells[col]

    @staticmethod
    def set_cell_size(t, row, col):
        t.rows[row].cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


WritingContract()
